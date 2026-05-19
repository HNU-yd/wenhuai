#!/usr/bin/env python3
"""Build a Word report for the add-project AHR mediation analysis."""

from __future__ import annotations

import math
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[1]
ADD_RESULTS = PROJECT / "results/add_project"
FIG_DIR = PROJECT / "results/figures"

OUTPUT_DOCX = ADD_RESULTS / "AHR_add_project_word_report.docx"

WALD_PATH = ADD_RESULTS / "AHR_two_significant_snp_wald_ratios.tsv"
STATUS_PATH = ADD_RESULTS / "AHR_two_snp_mediation_status.tsv"
COLOC_CONTEXT_PATH = ADD_RESULTS / "AHR_two_snp_coloc_posterior_context.tsv"
COLOC_MAIN_PATH = PROJECT / "results/coloc/AHR_eqtlgen_coloc_main_result.tsv"

FIGURES = [
    (
        "图 1. 中介 MR 分析框架",
        FIG_DIR / "AHR_two_snp_mediation_model.png",
    ),
    (
        "图 2. AHR cis-eQTL locus Manhattan plot（标注 rs17643734、rs59291726）",
        FIG_DIR / "AHR_eqtlgen_AHR_locus_manhattan_two_snp.png",
    ),
    (
        "图 3. 三套心肌炎 GWAS 的 AHR locus Manhattan plot",
        FIG_DIR / "AHR_myocarditis_locus_manhattan_two_snp.png",
    ),
    (
        "图 4. AHR eQTL 与心肌炎 GWAS 的 coloc 后验概率汇总",
        FIG_DIR / "AHR_coloc_summary_with_two_snp_context.png",
    ),
    (
        "图 5. 两个指定 SNP 的 AHR expression -> myocarditis Wald ratio forest plot",
        FIG_DIR / "AHR_two_snp_wald_ratio_forest.png",
    ),
]


def read_tsv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path, sep="\t", low_memory=False)


def is_missing(value: Any) -> bool:
    return value is None or pd.isna(value) or str(value) == ""


def num(value: Any) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return math.nan


def fmt(value: Any, digits: int = 3) -> str:
    if is_missing(value):
        return "NA"
    value_float = num(value)
    if not math.isfinite(value_float):
        return str(value)
    if value_float == 0:
        return "0"
    if abs(value_float) < 0.001:
        return f"{value_float:.{digits}e}"
    if abs(value_float) >= 1000:
        return f"{value_float:.{digits}e}"
    return f"{value_float:.{digits}f}"


def p_fmt(value: Any) -> str:
    return fmt(value, 3)


def or_ci(row: pd.Series) -> str:
    if is_missing(row.get("ratio_or")):
        return "NA"
    return (
        f"{fmt(row.get('ratio_or'))} "
        f"({fmt(row.get('ratio_or_lci95'))}-{fmt(row.get('ratio_or_uci95'))})"
    )


def interpret_wald(row: pd.Series) -> str:
    if row.get("status") != "ok":
        return "outcome 数据缺失，未计算"
    pval = num(row.get("ratio_pval"))
    ratio = num(row.get("ratio"))
    if not math.isfinite(pval):
        return "未得到有效 P 值"
    direction = "正向" if ratio > 0 else "负向"
    if pval < 0.05:
        return f"{direction}，达到 P<0.05"
    if pval < 0.10:
        return f"{direction}趋势，未达 P<0.05"
    return f"{direction}，不显著"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: Any, bold: bool = False) -> None:
    cell.text = "" if is_missing(text) else str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(8.5)
            run.bold = bold
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)

    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")


def add_paragraph(doc: Document, text: str = "", bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        text = text[len(bold_prefix) :]
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    styles["Normal"].font.size = Pt(10.5)

    return doc


def build_wald_summary(wald: pd.DataFrame) -> list[list[str]]:
    rows = []
    for _, row in wald.iterrows():
        rows.append(
            [
                row.get("outcome_dataset"),
                row.get("outcome_population"),
                row.get("target_snp"),
                row.get("status"),
                fmt(row.get("beta_exposure")),
                p_fmt(row.get("pval_exposure")),
                fmt(row.get("beta_outcome_aligned")),
                p_fmt(row.get("pval_outcome")),
                or_ci(row),
                p_fmt(row.get("ratio_pval")),
                interpret_wald(row),
            ]
        )
    return rows


def build_coloc_summary(coloc: pd.DataFrame) -> list[list[str]]:
    rows = []
    for _, row in coloc.iterrows():
        rows.append(
            [
                row.get("outcome_dataset"),
                fmt(row.get("n_overlap"), 0),
                fmt(row.get("PP.H3.abf")),
                fmt(row.get("PP.H4.abf")),
                row.get("coloc_interpretation"),
            ]
        )
    return rows


def build_snp_context(coloc_context: pd.DataFrame) -> list[list[str]]:
    rows = []
    for _, row in coloc_context.iterrows():
        rows.append(
            [
                row.get("outcome_dataset"),
                row.get("snp"),
                row.get("match_method"),
                fmt(row.get("beta_eqtl")),
                p_fmt(row.get("pval_eqtl")),
                fmt(row.get("beta_gwas_aligned")),
                p_fmt(row.get("pval_gwas")),
                fmt(row.get("SNP.PP.H4")),
            ]
        )
    return rows


def main() -> None:
    ADD_RESULTS.mkdir(parents=True, exist_ok=True)

    wald = read_tsv(WALD_PATH)
    status = read_tsv(STATUS_PATH)
    coloc_main = read_tsv(COLOC_MAIN_PATH)
    coloc_context = read_tsv(COLOC_CONTEXT_PATH)

    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AHR expression 与 Myocarditis 中介 MR 补充分析报告")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于加项目.docx 指定 SNP：rs17643734、rs59291726")
    run.font.size = Pt(10.5)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    add_paragraph(doc, "生成日期：2026-05-19")
    add_paragraph(doc, "项目目录：AHR_myocarditis_gwas")

    add_heading(doc, "一、分析目的", 1)
    add_paragraph(
        doc,
        "本报告根据加项目.docx 的要求，围绕 AHR expression GWAS "
        "eqtl-a-ENSG00000106546 和两个指定显著 SNP（rs17643734、rs59291726）"
        "补充中介 MR 相关分析、共定位结果、Manhattan/locus 图和结果解释。"
    )

    add_heading(doc, "二、分析流程", 1)
    add_bullets(
        doc,
        [
            "确定分析对象：AHR expression GWAS ID 为 eqtl-a-ENSG00000106546，指定 SNP 为 rs17643734 和 rs59291726。",
            "提取 exposure：从本地 eQTLGen AHR cis-eQTL 标准化结果中提取两个 SNP 的 AHR expression 效应值、标准误、P 值和等位基因信息。",
            "匹配 outcome：在 Sakaue2021 BBJ、Sakaue2021 EUR 和 FinnGen R12 I9_MYOCARD 三套心肌炎 GWAS 的 AHR locus 数据中查找对应 SNP。",
            "等位基因协调：按 exposure effect allele 对齐 outcome beta，必要时进行方向翻转；等位基因不匹配或 outcome 缺失则不计算 Wald ratio。",
            "计算单 SNP Wald ratio：使用 beta_outcome_aligned / beta_exposure 估计 AHR expression 对 myocarditis 的遗传预测效应，并用 delta method 计算标准误、P 值和 OR 及 95% CI。",
            "共定位分析：复用 AHR eQTL 与心肌炎 GWAS 的 coloc.abf 结果，汇总 PP.H3、PP.H4，并提取指定 SNP 的 SNP-level posterior context。",
            "可视化：生成中介模式图、AHR eQTL locus Manhattan 图、心肌炎 locus Manhattan 图、coloc 汇总图和 Wald ratio forest plot。",
        ],
    )

    add_heading(doc, "三、输入数据与指定 SNP", 1)
    add_table(
        doc,
        ["类别", "内容"],
        [
            ["Exposure", "AHR expression eQTLGen whole blood，eqtl-a-ENSG00000106546"],
            ["指定 SNP", "rs17643734；rs59291726"],
            ["Outcome 1", "Sakaue2021_BBJ_Myocarditis，East Asian / BBJ"],
            ["Outcome 2", "Sakaue2021_EUR_Myocarditis，European"],
            ["Outcome 3", "FinnGen_R12_I9_MYOCARD，Finnish / European"],
            ["主要方法", "单 SNP Wald ratio；coloc.abf；locus/Manhattan 可视化"],
        ],
    )

    add_heading(doc, "四、Wald Ratio 结果", 1)
    add_table(
        doc,
        [
            "Outcome",
            "人群",
            "SNP",
            "状态",
            "beta_eQTL",
            "P_eQTL",
            "beta_GWAS",
            "P_GWAS",
            "OR (95% CI)",
            "P_Wald",
            "解释",
        ],
        build_wald_summary(wald),
    )

    add_paragraph(
        doc,
        "结果解读：rs17643734 在 BBJ outcome 中缺失，无法计算 Wald ratio。"
        "rs59291726 在 BBJ 中 OR 点估计较大，但置信区间极宽且 P=0.379，证据不足。"
        "Sakaue EUR 中两个 SNP 均为正向，其中 rs59291726 呈趋势性正向结果（P=0.089），但未达到 P<0.05。"
        "FinnGen 中两个 SNP 的 OR 均接近 1，P 值接近 1，未见效应信号。",
    )

    add_heading(doc, "五、中介 MR 状态", 1)
    add_table(
        doc,
        ["分析模块", "状态", "方法", "SNP 数", "说明"],
        [
            [
                row.get("mediation_arm"),
                row.get("status"),
                row.get("method"),
                fmt(row.get("nsnp"), 0),
                row.get("note"),
            ]
            for _, row in status.iterrows()
        ],
    )

    add_heading(doc, "六、Coloc 结果", 1)
    add_table(
        doc,
        ["Outcome", "重叠 SNP 数", "PP.H3", "PP.H4", "解释"],
        build_coloc_summary(coloc_main),
    )
    add_paragraph(
        doc,
        "Coloc 结果解读：三套心肌炎 GWAS 的 PP.H4 分别为 0.062、0.048 和 0.021，"
        "均明显低于常用强共定位参考阈值 0.8。因此，AHR eQTL 与心肌炎 GWAS 在该区域缺乏强共定位证据。"
    )

    add_heading(doc, "七、指定 SNP 的 Coloc Posterior Context", 1)
    add_table(
        doc,
        ["Outcome", "SNP", "匹配方式", "beta_eQTL", "P_eQTL", "beta_GWAS", "P_GWAS", "SNP.PP.H4"],
        build_snp_context(coloc_context),
    )
    add_paragraph(
        doc,
        "注意：SNP.PP.H4 是在 H4 假设下的 SNP-level posterior context，不能替代整体 PP.H4。"
        "整体共定位判断仍应以每个 outcome 的 PP.H4.abf 为准。"
    )

    add_heading(doc, "八、图像结果", 1)
    for caption, figure_path in FIGURES:
        if not figure_path.exists():
            add_paragraph(doc, f"{caption}：图片缺失，路径 {figure_path}")
            continue
        doc.add_picture(str(figure_path), width=Inches(6.4))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, caption)

    add_heading(doc, "九、综合结论", 1)
    add_bullets(
        doc,
        [
            "两个指定 SNP 在 AHR expression GWAS 中均为强 eQTL，F 统计量较高，适合作为本次文档指定的 AHR expression 工具 SNP 进行补充分析。",
            "AHR expression -> myocarditis 的单 SNP Wald ratio 在三套心肌炎 GWAS 中均未达到统计显著。",
            "Sakaue EUR 中 rs59291726 呈正向趋势，但 P=0.089，不能作为明确阳性证据。",
            "FinnGen 中两个 SNP 的效应接近零，未支持 AHR expression 对 myocarditis 的遗传预测效应。",
            "Coloc 分析中三套数据 PP.H4 均较低，提示 AHR eQTL 与心肌炎 GWAS 信号缺乏强共定位支持。",
            "总体而言，当前结果不支持 AHR expression 作为心肌炎风险的明确遗传中介；该结论受限于指定 SNP 数量少、部分 outcome SNP 缺失以及心肌炎 GWAS 局部信号较弱。",
        ],
    )

    add_heading(doc, "十、输出文件", 1)
    add_bullets(
        doc,
        [
            "results/add_project/AHR_two_significant_snp_wald_ratios.tsv",
            "results/add_project/AHR_two_snp_mediation_status.tsv",
            "results/add_project/AHR_two_snp_coloc_posterior_context.tsv",
            "results/add_project/AHR_add_project_report.md",
            "results/add_project/AHR_add_project_word_report.docx",
            "results/figures/AHR_two_snp_mediation_model.png",
            "results/figures/AHR_eqtlgen_AHR_locus_manhattan_two_snp.png",
            "results/figures/AHR_myocarditis_locus_manhattan_two_snp.png",
            "results/figures/AHR_coloc_summary_with_two_snp_context.png",
            "results/figures/AHR_two_snp_wald_ratio_forest.png",
        ],
    )

    doc.save(OUTPUT_DOCX)
    print(f"Wrote {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
